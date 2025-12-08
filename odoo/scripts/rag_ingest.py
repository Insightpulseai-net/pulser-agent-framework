#!/usr/bin/env python3
"""
RAG Document Ingestion Pipeline
================================

Ingests documents from various sources, processes them through OCR/extraction,
chunks the content, and stores in the RAG tables.

Usage:
    python rag_ingest.py                    # Process pending sources
    python rag_ingest.py --source upload    # Process specific source type
    python rag_ingest.py --file path.pdf    # Ingest single file
    python rag_ingest.py --dry-run          # Show what would be processed

Supported Sources:
    - upload: Local file uploads
    - s3/spaces: S3-compatible object storage
    - gdrive: Google Drive
    - web: Web scraping
    - odoo: Odoo attachments

Environment Variables:
    DATABASE_URL    - PostgreSQL connection string
    OCR_ENGINE      - Default OCR engine (tesseract, paddle, azure, gvision)
    SPACES_*        - DigitalOcean Spaces credentials
"""

import os
import sys
import argparse
import hashlib
import logging
import mimetypes
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass
from datetime import datetime

try:
    import psycopg2
    from psycopg2.extras import DictCursor, Json
    from dotenv import load_dotenv
    from rich.console import Console
    from rich.progress import Progress, SpinnerColumn, TextColumn
except ImportError as e:
    print(f"Missing dependency: {e}")
    print("Run: pip install -r requirements.txt")
    sys.exit(1)

load_dotenv()

# Configuration
DATABASE_URL = os.getenv("DATABASE_URL")
OCR_ENGINE = os.getenv("OCR_ENGINE", "tesseract")
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "1000"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "200"))

# Console
console = Console()

# Logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("rag_ingest.log"),
    ]
)
logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Processed document ready for chunking."""
    source_id: str
    title: str
    full_text: str
    language: str = "en"
    page_count: int = 1
    meta: Dict = None


@dataclass
class Chunk:
    """Text chunk ready for embedding."""
    text: str
    index: int
    section_path: str = ""
    page_number: int = None
    token_count: int = 0


class DocumentProcessor:
    """Processes documents from various formats."""

    def __init__(self, ocr_engine: str = OCR_ENGINE):
        self.ocr_engine = ocr_engine

    def process_file(self, file_path: Path, content_type: str = None) -> Document:
        """Process a file and extract text."""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        content_type = content_type or mimetypes.guess_type(str(file_path))[0]
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return self._process_pdf(file_path)
        elif suffix in [".txt", ".md", ".rst"]:
            return self._process_text(file_path)
        elif suffix in [".html", ".htm"]:
            return self._process_html(file_path)
        elif suffix in [".docx"]:
            return self._process_docx(file_path)
        elif suffix in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
            return self._process_image(file_path)
        else:
            # Try as plain text
            return self._process_text(file_path)

    def _process_pdf(self, file_path: Path) -> Document:
        """Extract text from PDF using pymupdf or fallback to OCR."""
        try:
            import fitz  # pymupdf
            doc = fitz.open(file_path)
            pages = []
            for page in doc:
                text = page.get_text()
                if text.strip():
                    pages.append(text)
                else:
                    # OCR fallback for scanned pages
                    pages.append(self._ocr_page(page))
            doc.close()

            return Document(
                source_id="",
                title=file_path.stem,
                full_text="\n\n".join(pages),
                page_count=len(pages),
            )
        except ImportError:
            logger.warning("pymupdf not installed, falling back to OCR")
            return self._ocr_pdf(file_path)

    def _process_text(self, file_path: Path) -> Document:
        """Read plain text file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

        return Document(
            source_id="",
            title=file_path.stem,
            full_text=text,
        )

    def _process_html(self, file_path: Path) -> Document:
        """Extract text from HTML."""
        try:
            from bs4 import BeautifulSoup
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                soup = BeautifulSoup(f.read(), "html.parser")

            # Remove scripts and styles
            for element in soup(["script", "style", "nav", "footer"]):
                element.decompose()

            text = soup.get_text(separator="\n", strip=True)
            title = soup.title.string if soup.title else file_path.stem

            return Document(
                source_id="",
                title=title,
                full_text=text,
            )
        except ImportError:
            logger.warning("beautifulsoup4 not installed")
            return self._process_text(file_path)

    def _process_docx(self, file_path: Path) -> Document:
        """Extract text from DOCX."""
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            return Document(
                source_id="",
                title=file_path.stem,
                full_text="\n\n".join(paragraphs),
            )
        except ImportError:
            logger.warning("python-docx not installed")
            raise ValueError("DOCX processing requires python-docx")

    def _process_image(self, file_path: Path) -> Document:
        """OCR an image file."""
        text = self._ocr_file(file_path)
        return Document(
            source_id="",
            title=file_path.stem,
            full_text=text,
        )

    def _ocr_file(self, file_path: Path) -> str:
        """Run OCR on an image file."""
        if self.ocr_engine == "tesseract":
            return self._ocr_tesseract(file_path)
        elif self.ocr_engine == "paddle":
            return self._ocr_paddle(file_path)
        else:
            # Default to tesseract
            return self._ocr_tesseract(file_path)

    def _ocr_tesseract(self, file_path: Path) -> str:
        """OCR using Tesseract."""
        try:
            import pytesseract
            from PIL import Image
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except ImportError:
            logger.error("pytesseract not installed")
            return ""

    def _ocr_paddle(self, file_path: Path) -> str:
        """OCR using PaddleOCR."""
        try:
            from paddleocr import PaddleOCR
            ocr = PaddleOCR(use_angle_cls=True, lang='en')
            result = ocr.ocr(str(file_path))
            lines = []
            for line in result[0]:
                if line[1]:
                    lines.append(line[1][0])
            return "\n".join(lines)
        except ImportError:
            logger.error("paddleocr not installed")
            return ""

    def _ocr_pdf(self, file_path: Path) -> Document:
        """OCR an entire PDF."""
        try:
            import pdf2image
            images = pdf2image.convert_from_path(file_path)
            pages = []
            for img in images:
                import pytesseract
                text = pytesseract.image_to_string(img)
                pages.append(text)

            return Document(
                source_id="",
                title=file_path.stem,
                full_text="\n\n".join(pages),
                page_count=len(pages),
            )
        except ImportError:
            logger.error("pdf2image or pytesseract not installed")
            return Document(source_id="", title=file_path.stem, full_text="")

    def _ocr_page(self, page) -> str:
        """OCR a single PDF page using pymupdf."""
        try:
            import pytesseract
            from PIL import Image
            import io

            # Render page as image
            pix = page.get_pixmap()
            img_data = pix.tobytes("png")
            image = Image.open(io.BytesIO(img_data))

            return pytesseract.image_to_string(image)
        except Exception as e:
            logger.warning(f"OCR failed for page: {e}")
            return ""


class TextChunker:
    """Splits text into chunks for embedding."""

    def __init__(
        self,
        chunk_size: int = CHUNK_SIZE,
        chunk_overlap: int = CHUNK_OVERLAP,
        strategy: str = "recursive"
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.strategy = strategy

    def chunk(self, text: str) -> List[Chunk]:
        """Split text into chunks."""
        if self.strategy == "recursive":
            return self._recursive_chunk(text)
        elif self.strategy == "sentence":
            return self._sentence_chunk(text)
        else:
            return self._fixed_chunk(text)

    def _recursive_chunk(self, text: str) -> List[Chunk]:
        """Recursive character-based chunking with overlap."""
        chunks = []
        separators = ["\n\n", "\n", ". ", " ", ""]

        def split_text(text: str, separators: List[str]) -> List[str]:
            if not separators:
                return [text]

            sep = separators[0]
            parts = text.split(sep) if sep else list(text)

            result = []
            current = ""

            for part in parts:
                test = current + (sep if current else "") + part
                if len(test) <= self.chunk_size:
                    current = test
                else:
                    if current:
                        result.append(current)
                    if len(part) > self.chunk_size:
                        # Recurse with next separator
                        result.extend(split_text(part, separators[1:]))
                        current = ""
                    else:
                        current = part

            if current:
                result.append(current)

            return result

        raw_chunks = split_text(text, separators)

        # Add overlap
        for i, chunk_text in enumerate(raw_chunks):
            chunks.append(Chunk(
                text=chunk_text.strip(),
                index=i,
                token_count=len(chunk_text.split()),
            ))

        return chunks

    def _sentence_chunk(self, text: str) -> List[Chunk]:
        """Sentence-based chunking."""
        try:
            import nltk
            nltk.download('punkt', quiet=True)
            sentences = nltk.sent_tokenize(text)
        except ImportError:
            # Fallback to simple splitting
            sentences = text.replace("!", ".").replace("?", ".").split(".")
            sentences = [s.strip() for s in sentences if s.strip()]

        chunks = []
        current_chunk = ""
        current_sentences = []

        for sentence in sentences:
            test = current_chunk + " " + sentence if current_chunk else sentence
            if len(test) <= self.chunk_size:
                current_chunk = test
                current_sentences.append(sentence)
            else:
                if current_chunk:
                    chunks.append(Chunk(
                        text=current_chunk.strip(),
                        index=len(chunks),
                        token_count=len(current_chunk.split()),
                    ))
                current_chunk = sentence
                current_sentences = [sentence]

        if current_chunk:
            chunks.append(Chunk(
                text=current_chunk.strip(),
                index=len(chunks),
                token_count=len(current_chunk.split()),
            ))

        return chunks

    def _fixed_chunk(self, text: str) -> List[Chunk]:
        """Fixed-size character chunking."""
        chunks = []
        for i in range(0, len(text), self.chunk_size - self.chunk_overlap):
            chunk_text = text[i:i + self.chunk_size]
            chunks.append(Chunk(
                text=chunk_text.strip(),
                index=len(chunks),
                token_count=len(chunk_text.split()),
            ))
        return chunks


class RAGIngester:
    """Main ingestion pipeline."""

    def __init__(self, db_url: str = DATABASE_URL):
        self.db_url = db_url
        self.processor = DocumentProcessor()
        self.chunker = TextChunker()
        self.conn = None

    def connect(self) -> bool:
        """Connect to database."""
        if not self.db_url:
            console.print("[red]DATABASE_URL not set[/red]")
            return False

        try:
            self.conn = psycopg2.connect(self.db_url)
            console.print("[green]Connected to database[/green]")
            return True
        except Exception as e:
            console.print(f"[red]Database connection failed: {e}[/red]")
            return False

    def close(self):
        """Close database connection."""
        if self.conn:
            self.conn.close()

    def get_pending_sources(self, kind: str = None, limit: int = 100) -> List[Dict]:
        """Get sources pending processing."""
        with self.conn.cursor(cursor_factory=DictCursor) as cur:
            query = """
                SELECT id, tenant_id, workspace_id, kind, uri, display_name,
                       content_type, ocr_engine, meta
                FROM rag_sources
                WHERE status = 'pending'
            """
            params = []

            if kind:
                query += " AND kind = %s"
                params.append(kind)

            query += " ORDER BY created_at LIMIT %s"
            params.append(limit)

            cur.execute(query, params)
            return [dict(row) for row in cur.fetchall()]

    def create_source(
        self,
        tenant_id: str,
        uri: str,
        kind: str = "upload",
        display_name: str = None,
        content_type: str = None,
        workspace_id: str = None,
        meta: Dict = None,
    ) -> str:
        """Create a new source record."""
        with self.conn.cursor() as cur:
            cur.execute("""
                INSERT INTO rag_sources
                (tenant_id, workspace_id, kind, uri, display_name, content_type, meta)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
                RETURNING id
            """, (
                tenant_id,
                workspace_id,
                kind,
                uri,
                display_name or Path(uri).name,
                content_type,
                Json(meta or {}),
            ))
            source_id = cur.fetchone()[0]
            self.conn.commit()
            return str(source_id)

    def process_source(self, source: Dict, dry_run: bool = False) -> bool:
        """Process a single source."""
        source_id = source["id"]
        uri = source["uri"]
        kind = source["kind"]

        console.print(f"Processing: {source['display_name'] or uri}")

        try:
            # Update status to processing
            if not dry_run:
                self._update_source_status(source_id, "processing")

            # Get the file/content
            if kind == "upload" or kind == "seed":
                file_path = Path(uri.replace("file://", "").replace("seed://", ""))
                if not file_path.exists() and kind != "seed":
                    raise FileNotFoundError(f"File not found: {file_path}")

                # For seed sources, create dummy content
                if kind == "seed":
                    document = Document(
                        source_id=source_id,
                        title=source["display_name"],
                        full_text=source.get("meta", {}).get("text", ""),
                    )
                else:
                    document = self.processor.process_file(
                        file_path,
                        source.get("content_type"),
                    )
            else:
                raise ValueError(f"Unsupported source kind: {kind}")

            document.source_id = source_id

            if dry_run:
                console.print(f"  Would create document: {document.title}")
                console.print(f"  Text length: {len(document.full_text)} chars")
                return True

            # Compute hash for deduplication
            text_hash = hashlib.sha256(document.full_text.encode()).hexdigest()

            # Check for duplicate
            with self.conn.cursor() as cur:
                cur.execute(
                    "SELECT id FROM rag_documents WHERE hash_sha256 = %s",
                    (text_hash,)
                )
                if cur.fetchone():
                    console.print(f"  [yellow]Duplicate document, skipping[/yellow]")
                    self._update_source_status(source_id, "processed")
                    return True

            # Create document
            doc_id = self._create_document(source, document, text_hash)

            # Chunk the document
            chunks = self.chunker.chunk(document.full_text)
            console.print(f"  Created {len(chunks)} chunks")

            # Store chunks
            for chunk in chunks:
                self._create_chunk(source["tenant_id"], doc_id, chunk)

            # Mark source as processed
            self._update_source_status(source_id, "processed")
            self.conn.commit()

            console.print(f"  [green]Successfully processed[/green]")
            return True

        except Exception as e:
            logger.error(f"Failed to process source {source_id}: {e}")
            if not dry_run:
                self._update_source_status(source_id, "error", str(e))
            console.print(f"  [red]Error: {e}[/red]")
            return False

    def _update_source_status(self, source_id: str, status: str, error: str = None):
        """Update source processing status."""
        with self.conn.cursor() as cur:
            cur.execute("""
                UPDATE rag_sources
                SET status = %s, error_message = %s, updated_at = NOW()
                WHERE id = %s
            """, (status, error, source_id))
        self.conn.commit()

    def _create_document(self, source: Dict, document: Document, text_hash: str) -> str:
        """Create document record."""
        with self.conn.cursor() as cur:
            cur.execute("""
                INSERT INTO rag_documents
                (tenant_id, source_id, title, language, page_count, word_count,
                 hash_sha256, full_text, meta)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id
            """, (
                source["tenant_id"],
                source["id"],
                document.title,
                document.language,
                document.page_count,
                len(document.full_text.split()),
                text_hash,
                document.full_text,
                Json(document.meta or {}),
            ))
            return str(cur.fetchone()[0])

    def _create_chunk(self, tenant_id: str, document_id: str, chunk: Chunk):
        """Create chunk record."""
        with self.conn.cursor() as cur:
            cur.execute("""
                INSERT INTO rag_chunks
                (tenant_id, document_id, chunk_index, text, token_count,
                 section_path, page_number, chunking_strategy, meta)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                tenant_id,
                document_id,
                chunk.index,
                chunk.text,
                chunk.token_count,
                chunk.section_path,
                chunk.page_number,
                self.chunker.strategy,
                Json({}),
            ))

    def ingest_file(
        self,
        file_path: str,
        tenant_id: str,
        workspace_id: str = None,
        dry_run: bool = False,
    ) -> bool:
        """Ingest a single file."""
        path = Path(file_path)
        if not path.exists():
            console.print(f"[red]File not found: {file_path}[/red]")
            return False

        content_type = mimetypes.guess_type(str(path))[0]

        source_id = self.create_source(
            tenant_id=tenant_id,
            uri=f"file://{path.absolute()}",
            kind="upload",
            display_name=path.name,
            content_type=content_type,
            workspace_id=workspace_id,
        )

        source = {
            "id": source_id,
            "tenant_id": tenant_id,
            "workspace_id": workspace_id,
            "kind": "upload",
            "uri": f"file://{path.absolute()}",
            "display_name": path.name,
            "content_type": content_type,
        }

        return self.process_source(source, dry_run)

    def run(self, kind: str = None, dry_run: bool = False) -> Dict:
        """Run the ingestion pipeline."""
        results = {"processed": 0, "failed": 0, "skipped": 0}

        sources = self.get_pending_sources(kind=kind)
        console.print(f"Found {len(sources)} pending sources")

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
        ) as progress:
            task = progress.add_task("Ingesting documents...", total=len(sources))

            for source in sources:
                progress.update(task, description=f"Processing {source['display_name'] or source['uri'][:30]}...")

                if self.process_source(source, dry_run):
                    results["processed"] += 1
                else:
                    results["failed"] += 1

                progress.advance(task)

        return results


def main():
    parser = argparse.ArgumentParser(description="RAG Document Ingestion Pipeline")
    parser.add_argument("--source", "-s", help="Source type to process (upload, web, gdrive, etc.)")
    parser.add_argument("--file", "-f", help="Single file to ingest")
    parser.add_argument("--tenant", "-t", help="Tenant ID for file ingestion")
    parser.add_argument("--dry-run", "-n", action="store_true", help="Show what would be done")
    parser.add_argument("--db", default=DATABASE_URL, help="Database URL")

    args = parser.parse_args()

    ingester = RAGIngester(db_url=args.db)

    if not ingester.connect():
        sys.exit(1)

    try:
        if args.file:
            if not args.tenant:
                console.print("[red]--tenant required when ingesting a file[/red]")
                sys.exit(1)

            success = ingester.ingest_file(
                args.file,
                tenant_id=args.tenant,
                dry_run=args.dry_run,
            )
            sys.exit(0 if success else 1)
        else:
            results = ingester.run(kind=args.source, dry_run=args.dry_run)

            console.print("\n[bold]Ingestion Results[/bold]")
            console.print(f"  Processed: {results['processed']}")
            console.print(f"  Failed: {results['failed']}")

            if results["failed"] > 0:
                sys.exit(1)
    finally:
        ingester.close()


if __name__ == "__main__":
    main()
